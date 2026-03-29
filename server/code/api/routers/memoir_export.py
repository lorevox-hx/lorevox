"""
Lorevox Memoir Export Router
=============================
Provides server-side DOCX export for memoir content.

Endpoints:
  POST  /api/memoir/export-docx  — accept memoir JSON, return .docx file

Design rules:
  - Export reflects exactly what the user sees (threads or draft).
  - Scaffold placeholder content is never exported.
  - Meaning sections (Turning Points, Hard Moments, etc.) become DOCX headings.
  - Structural sections (Family & Relationships, Work, etc.) become secondary headings.
  - Draft state is rendered as plain prose paragraphs with section headers.
  - Threads state is rendered as grouped bullet lists per section.
  - Media Builder: attached_photos inlines images after section headings (graceful skip on error).
"""

from __future__ import annotations

import io
import logging
from pathlib import Path
from typing import Any, Dict, List, Optional

from fastapi import APIRouter
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field

logger = logging.getLogger("memoir_export")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

router = APIRouter(prefix="/api/memoir", tags=["memoir-export"])


# ── Request models ─────────────────────────────────────────────────────────────

class MemoirSection(BaseModel):
    """A single named section with zero or more thread items."""
    id: str
    label: str
    items: List[str] = Field(default_factory=list)


class AttachedPhoto(BaseModel):
    """A photo attached to a memoir section (Media Builder — Task 4)."""
    media_id: str
    section_key: str
    file_path: str          # absolute server-local path; python-docx reads this directly
    description: str = ""
    taken_at: str = ""


class MemoirExportRequest(BaseModel):
    """
    Shape sent by the frontend's memoirExportDOCX() function.
    memoir_state: "threads" | "draft"
    narrator_name: display name for the document title
    sections: populated sections only (empty sections are pre-filtered by the caller)
    prose: flat prose string for draft state (paragraphs joined by \\n\\n)
    arc_roles: which narrative arc parts are present (display only, optional)
    attached_photos: photos to inline at their section (empty list = no change in behavior)
    """
    narrator_name: str = Field(default="Narrator")
    memoir_state: str = Field(default="threads")
    sections: List[MemoirSection] = Field(default_factory=list)
    prose: Optional[str] = Field(default=None)
    arc_roles: List[str] = Field(default_factory=list)
    attached_photos: List[AttachedPhoto] = Field(default_factory=list)


# ── Helpers ────────────────────────────────────────────────────────────────────

# Colour constants for the Lorevox brand tone (dark warm palette)
# Guarded: RGBColor only exists when python-docx is installed.
if _DOCX_AVAILABLE:
    _DARK_BROWN = RGBColor(0x3B, 0x2A, 0x1A)   # heading primary
    _WARM_GREY  = RGBColor(0x5A, 0x55, 0x50)   # heading secondary
    _GOLD       = RGBColor(0xAA, 0x88, 0x44)   # accent line / arc label
else:
    _DARK_BROWN = _WARM_GREY = _GOLD = None


def _photos_for_section(req: MemoirExportRequest, section_key: str) -> List[AttachedPhoto]:
    """Return all photos attached to a given memoir section key."""
    return [p for p in req.attached_photos if p.section_key == section_key]


def _add_photo_to_doc(doc: Any, photo: AttachedPhoto) -> None:
    """
    Insert photo inline in the document.
    Gracefully skips on any error (corrupt file, format unsupported, missing file).
    """
    try:
        path = Path(photo.file_path)
        if not path.exists():
            logger.warning("[memoir-docx] Photo not found on disk: %s — skipping", path)
            return
        doc.add_picture(str(path), width=Inches(3.5))
        # Caption paragraph
        caption_parts = []
        if photo.description:
            caption_parts.append(photo.description)
        if photo.taken_at:
            caption_parts.append(photo.taken_at)
        if caption_parts:
            cap = doc.add_paragraph(" — ".join(caption_parts))
            if cap.runs:
                cap.runs[0].font.size = Pt(9)
                cap.runs[0].font.italic = True
                cap.runs[0].font.color.rgb = _WARM_GREY
    except Exception as exc:
        logger.warning("[memoir-docx] Could not add photo %s: %s — skipping", photo.file_path, exc)


# ── DOCX builders ──────────────────────────────────────────────────────────────

def _build_threads_docx(req: MemoirExportRequest) -> bytes:
    """Build DOCX for threads state: grouped sections with bullet items."""
    doc = Document()

    # Title
    title = doc.add_heading(f"Story Threads — {req.narrator_name}", level=0)
    title.runs[0].font.color.rgb = _DARK_BROWN

    # Subtitle
    sub = doc.add_paragraph("Building Blocks Collected")
    sub.runs[0].font.italic = True
    sub.runs[0].font.color.rgb = _WARM_GREY

    # Arc coverage line (if available)
    if req.arc_roles:
        arc_line = doc.add_paragraph()
        arc_run = arc_line.add_run(f"Story arc: {' · '.join(req.arc_roles)}")
        arc_run.font.size = Pt(10)
        arc_run.font.color.rgb = _GOLD

    doc.add_paragraph()  # spacer

    # Sections
    for sec in req.sections:
        if not sec.items:
            continue  # skip empty — export truth rule
        h = doc.add_heading(sec.label, level=2)
        h.runs[0].font.color.rgb = _DARK_BROWN

        # Inline photos for this section (Media Builder)
        for photo in _photos_for_section(req, sec.id):
            _add_photo_to_doc(doc, photo)

        for item in sec.items:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(item)

        doc.add_paragraph()  # spacer between sections

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_draft_docx(req: MemoirExportRequest) -> bytes:
    """Build DOCX for draft state: prose paragraphs, optionally with arc headings."""
    doc = Document()

    title = doc.add_heading(f"Memoir Draft — {req.narrator_name}", level=0)
    title.runs[0].font.color.rgb = _DARK_BROWN

    sub = doc.add_paragraph("Your Words, Taking Shape")
    sub.runs[0].font.italic = True
    sub.runs[0].font.color.rgb = _WARM_GREY

    doc.add_paragraph()  # spacer

    # Build a map of section_key → photos for quick lookup in arc-label detection
    # We use section keys stored on the photo; for draft, we try to match arc labels
    # to memoir section ids (best-effort — draft state doesn't have structured sections).
    section_photos_by_key: dict = {}
    for photo in req.attached_photos:
        section_photos_by_key.setdefault(photo.section_key, []).append(photo)

    if req.prose:
        paragraphs = [p.strip() for p in req.prose.split("\n\n") if p.strip()]
        for para_text in paragraphs:
            lines = para_text.split("\n")
            # Detect arc label marker: "-- Label --"
            if lines and lines[0].strip().startswith("--") and lines[0].strip().endswith("--"):
                label = lines[0].strip().strip("-").strip()
                h = doc.add_heading(label, level=2)
                h.runs[0].font.color.rgb = _DARK_BROWN
                body = "\n".join(lines[1:]).strip()
                if body:
                    doc.add_paragraph(body)
            else:
                doc.add_paragraph(para_text)
            doc.add_paragraph()  # spacer

    # Append photo section at end of draft (no per-section matching in pure prose)
    # Only include photos not already displayed via section matching
    if req.attached_photos:
        doc.add_page_break()
        ph = doc.add_heading("Photos", level=1)
        ph.runs[0].font.color.rgb = _DARK_BROWN
        for photo in req.attached_photos:
            _add_photo_to_doc(doc, photo)
            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Route ─────────────────────────────────────────────────────────────────────

@router.post("/export-docx")
def api_memoir_export_docx(req: MemoirExportRequest):
    """
    Accept memoir content JSON, return a DOCX file as a streaming download.
    Called by memoirExportDOCX() in lori8.0.html.
    """
    if not _DOCX_AVAILABLE:
        from fastapi import HTTPException
        raise HTTPException(
            status_code=503,
            detail="python-docx is not installed on this server. Install with: pip install python-docx",
        )

    safe_name = (
        req.narrator_name.strip().lower()
        .replace(" ", "_")
        .replace("/", "_")
    )[:40] or "narrator"
    filename = f"lorevox_memoir_{safe_name}_{req.memoir_state}.docx"

    if req.memoir_state == "draft":
        docx_bytes = _build_draft_docx(req)
    else:
        docx_bytes = _build_threads_docx(req)

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
